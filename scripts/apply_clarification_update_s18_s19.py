#!/usr/bin/env python3
"""Clarification update: multilabel S18, reachability S19, GAT note, S7 wording, WC."""
from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

ROOT = Path(__file__).resolve().parents[1]
MAIN = ROOT / "01_Main_Manuscript.docx"
SUPP = ROOT / "02_Supplementary_Material.docx"
COVER_TXT = ROOT / "00_Cover_Letter.txt"
COVER_DOCX = ROOT / "00_Cover_Letter.docx"
MULTI = ROOT / "analysis" / "multilabel_any_known_positive.json"
REACH = ROOT / "analysis" / "gallery_construction_reachability.json"

LINE, RULE = "480", "auto"


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = text


def ensure_pPr(p):
    pPr = p._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = etree.Element(qn("w:pPr"))
        p._element.insert(0, pPr)
    return pPr


def set_line(p):
    pPr = ensure_pPr(p)
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = etree.SubElement(pPr, qn("w:spacing"))
    sp.set(qn("w:line"), LINE)
    sp.set(qn("w:lineRule"), RULE)


def set_run_tnr(p, half_pt="22"):
    for run in p.runs:
        rPr = run._element.find(qn("w:rPr"))
        if rPr is None:
            rPr = etree.SubElement(run._element, qn("w:rPr"))
            run._element.insert(0, rPr)
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = etree.SubElement(rPr, qn("w:rFonts"))
            rPr.insert(0, rFonts)
        for a in (qn("w:asciiTheme"), qn("w:hAnsiTheme"), qn("w:eastAsiaTheme"), qn("w:cstheme")):
            if a in rFonts.attrib:
                del rFonts.attrib[a]
        for a in (qn("w:ascii"), qn("w:hAnsi"), qn("w:cs"), qn("w:eastAsia")):
            rFonts.set(a, "Times New Roman")
        for tag in ("w:sz", "w:szCs"):
            el = rPr.find(qn(tag))
            if el is None:
                el = etree.SubElement(rPr, qn(tag))
            el.set(qn("w:val"), half_pt)


def insert_before(paragraph, text: str, style_name: str):
    new_p = paragraph._element.makeelement(qn("w:p"), {})
    paragraph._element.addprevious(new_p)
    new_para = type(paragraph)(new_p, paragraph._parent)
    new_para.style = style_name
    new_para.add_run(text)
    set_line(new_para)
    set_run_tnr(new_para, "22" if not style_name.startswith("Heading") else "22")
    return new_para


def pct(x: float, d: int = 1) -> str:
    return f"{100 * x:.{d}f}%"


def word_counts(doc: Document) -> tuple[int, int]:
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    body = paras[paras.index("Introduction") : paras.index("Declarations")]
    with_legends = " ".join(body)
    without = " ".join(
        p
        for p in body
        if not re.match(r"^(Figure|Table) \d+\.", p) and not p.startswith("Note.")
    )
    table_text = " ".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    return len((with_legends + " " + table_text).split()), len(
        (without + " " + table_text).split()
    )


def update_cover(wc_with: int, wc_without: int) -> None:
    scope_pat = (
        r"The main text from Introduction through (?:Conclusion|Declarations) is [\d,]+ words "
        r"including in-text figure/table legends \([\d,]+ words excluding those legends\), "
        r"excluding Abstract, References and Supplementary Material\."
    )
    repl = (
        f"The main text from Introduction through Declarations is {wc_with:,} words "
        f"including in-text figure/table legends ({wc_without:,} words excluding those legends), "
        "excluding Abstract, References and Supplementary Material."
    )
    # clarification sentence for resubmission comment / cover
    clar = (
        "This transfer package is a clarification and reproducibility update: it clarifies "
        "designated-positive versus any-known-positive evaluation under incomplete multi-label "
        "records, documents gallery construction and operational reachability, and aligns "
        "manuscript text with deposited code/archives. It does not change the primary "
        "fixed-gallery diagnostic conclusions."
    )
    text = COVER_TXT.read_text()
    text2, n = re.subn(scope_pat, repl, text)
    if n == 0:
        raise SystemExit("cover WC sentence missing")
    if "clarification and reproducibility update" not in text2:
        # insert before closing thanks / sincerely if present
        if "Yours sincerely" in text2:
            text2 = text2.replace("Yours sincerely", clar + "\n\nYours sincerely")
        elif "Sincerely" in text2:
            text2 = text2.replace("Sincerely", clar + "\n\nSincerely")
        else:
            text2 = text2.rstrip() + "\n\n" + clar + "\n"
    COVER_TXT.write_text(text2)

    doc = Document(str(COVER_DOCX))
    hit_wc = hit_clar = False
    for p in doc.paragraphs:
        if re.search(scope_pat, p.text):
            set_paragraph_text(p, re.sub(scope_pat, repl, p.text))
            hit_wc = True
    if "clarification and reproducibility update" not in "\n".join(
        p.text for p in doc.paragraphs
    ):
        # append paragraph before end
        doc.add_paragraph(clar)
        hit_clar = True
    if not hit_wc:
        raise SystemExit("cover docx WC missing")
    doc.save(str(COVER_DOCX))
    print("cover updated", "clar_added" if hit_clar else "clar_present")


def patch_main(doc: Document, multi: dict, reach: dict) -> None:
    h = multi["hvidb_test"]
    i = multi["intact_cross_test"]
    hs = h["stage1_retrieval"]
    hf = h["fusion_T20"]

    # Methods: query definition
    for p in doc.paragraphs:
        if p.text.startswith("Each query is a human protein"):
            if "designated-positive retrieval under incomplete multi-label" not in p.text:
                set_paragraph_text(
                    p,
                    p.text.rstrip(".")
                    + ". We treat this as designated-positive retrieval under incomplete "
                    "multi-label interaction records: each query has one designated true "
                    "partner v⁺, while the same human may have additional evaluation-file "
                    "positive partners that are not scored as the designated label. The "
                    "primary Hit@K / Recall@T / oracle@T identity uses the designated "
                    "partner only; any-known-positive Hit@K is reported separately as a "
                    "multilabel sensitivity (Supplementary Table S18).",
                )
            break

    # Methods: GAT clarification in two-stage setup
    for p in doc.paragraphs:
        if p.text.startswith("Analyses use a two-stage retrieve-then-rerank pipeline"):
            if "training-time auxiliary module" not in p.text and "batch GAT" in p.text:
                set_paragraph_text(
                    p,
                    p.text.replace(
                        "At inference the batch graph and GAT are omitted: each gallery virus is encoded independently via frozen ESM3 and the projection head, then L2-normalized for exhaustive cosine retrieval.",
                        "The batch GAT is a training-time auxiliary module on positive-pair edges within each training batch; at inference there is no batch graph, no edge construction and no dependence on test labels. "
                        "Each query and gallery protein is encoded independently via frozen ESM3 and the projection head, then L2-normalized for exhaustive cosine retrieval.",
                    ),
                )
            elif "training-time auxiliary module" not in p.text:
                # insert after GAT mention if present differently
                if "GAT" in p.text and "At inference" in p.text:
                    set_paragraph_text(
                        p,
                        p.text.replace(
                            "At inference",
                            "The batch GAT is used only during training as an auxiliary module; "
                            "at inference",
                        ),
                    )
            break

    # Methods: reachability operational note in sources of protocol variation / gallery para
    for p in doc.paragraphs:
        if "reachability is 100% by construction" in p.text and "operational property" not in p.text:
            set_paragraph_text(
                p,
                p.text
                + " Thus 100% reachability is an operational property of gallery construction "
                "and filtering, not empirical retrieval success and not the natural coverage "
                "rate of incomplete databases (Supplementary Table S19).",
            )
            break

    # Table 2 note
    for p in doc.paragraphs:
        if p.text.startswith("Note. Ranking instantiations share the same gallery"):
            if "operational property of gallery construction" not in p.text:
                set_paragraph_text(
                    p,
                    p.text
                    + " Reachability was 100% by evaluation-gallery construction/filtering "
                    "(operational, not natural database coverage; Supplementary Table S19).",
                )
            # avoid duplicate if already said earlier in same note
            # earlier patch already had "Reachability was 100% by evaluation-gallery construction"
            break

    # Results: multilabel paragraph after overview / end-to-end
    multi_para = (
        f"Because many humans have multiple evaluation-file positive partners "
        f"({pct(h['frac_queries_multi_positive'])} of HVIDB test queries; "
        f"{pct(i['frac_queries_multi_positive'])} of IntAct), designated-positive Hit@K can "
        f"mark a query as a miss even when another known positive of the same human appears "
        f"in the top-K. Under Stage-1 retrieval on HVIDB test, designated-positive Hit@10 was "
        f"{pct(hs['designated_hit@10'])}, whereas any-known-positive Hit@10 rose to "
        f"{pct(hs['any_known_positive_hit@10'])}; fusion Hit@10 similarly rose from "
        f"{pct(hf['designated_hit@10'])} to {pct(hf['any_known_positive_hit@10'])} "
        f"(Supplementary Table S18). Any-known-positive metrics are a multilabel sensitivity "
        f"only; they do not replace the designated-positive protocol or the "
        f"Hit@K = Recall@T × oracle@K identity."
    )
    inserted = any("any-known-positive Hit@10" in p.text for p in doc.paragraphs)
    if not inserted:
        for p in doc.paragraphs:
            if p.text.startswith("On held-out HVIDB test, Stage-1 Recall@20 was"):
                # insert after this paragraph
                nxt = p._element.getnext()
                # create after current by inserting before next sibling
                # find paragraph object for next
                # simpler: append text to end of this para
                set_paragraph_text(p, p.text.rstrip() + " " + multi_para)
                break

    # S7 wording in limitations
    for p in doc.paragraphs:
        if "purposive illustrative snapshot" in p.text:
            set_paragraph_text(
                p, p.text.replace("purposive illustrative snapshot", "illustrative audit")
            )
        if "rather than a systematic review" in p.text and "illustrative audit" in p.text:
            # already fine
            pass


def append_supp(doc: Document, multi: dict, reach: dict) -> None:
    if any("Supplementary Table S18" in p.text for p in doc.paragraphs):
        return

    def add(text, style="Normal"):
        p = doc.add_paragraph(text)
        try:
            p.style = style
        except Exception:
            pass
        set_line(p)
        set_run_tnr(p)
        return p

    h = multi["hvidb_test"]
    irow = multi["intact_cross_test"]
    hs, is_ = h["stage1_retrieval"], irow["stage1_retrieval"]
    hf, if_ = h["fusion_T20"], irow["fusion_T20"]

    add(
        "Supplementary Table S18. Designated-positive versus any-known-positive Hit@K "
        "under incomplete multi-label records (fixed gallery)."
    )
    tbl = doc.add_table(rows=1, cols=9)
    headers = [
        "Dataset",
        "n / humans",
        "Multi-pos. queries",
        "Des Hit@1/5/10 (Stage-1)",
        "Any Hit@1/5/10 (Stage-1)",
        "Des Hit@10 (fusion)",
        "Any Hit@10 (fusion)",
        "Des Hit@10 (direct)",
        "Any Hit@10 (direct)",
    ]
    for j, htxt in enumerate(headers):
        tbl.rows[0].cells[j].text = htxt

    def row_vals(block, s1, fus, direct):
        return [
            block["label"],
            f"{block['n_queries']} / {block['n_unique_humans']}",
            pct(block["frac_queries_multi_positive"]),
            f"{pct(s1['designated_hit@1'])} / {pct(s1['designated_hit@5'])} / {pct(s1['designated_hit@10'])}",
            f"{pct(s1['any_known_positive_hit@1'])} / {pct(s1['any_known_positive_hit@5'])} / {pct(s1['any_known_positive_hit@10'])}",
            pct(fus["designated_hit@10"]),
            pct(fus["any_known_positive_hit@10"]),
            pct(direct["designated_hit@10"]),
            pct(direct["any_known_positive_hit@10"]),
        ]

    for block, s1, fus in (
        (h, hs, hf),
        (irow, is_, if_),
    ):
        r = tbl.add_row()
        for j, v in enumerate(row_vals(block, s1, fus, block["direct_full_gallery"])):
            r.cells[j].text = v
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                set_line(p)
                set_run_tnr(p, "18")

    add(
        "Note. Known positives = evaluation-file positive partners of the query human. "
        "Designated-positive is the primary protocol (Table 2; Hit@K = Recall@T × oracle@K). "
        "Any-known-positive Hit@K = 1 if any known positive is retrieved into the top-K "
        "(Stage-1) or achieves Hit@10 under fusion/direct for that partner. "
        f"HVIDB: median known positives per query = {h['median_known_positives_per_query']:.0f}. "
        "Any-known-positive metrics are multilabel sensitivities and are not interchangeable "
        "with the designated-positive identity."
    )

    add(
        "Supplementary Table S19. Gallery construction and operational reachability."
    )
    tbl2 = doc.add_table(rows=1, cols=7)
    h2 = [
        "Dataset",
        "Nominal IDs",
        "Sequence-available",
        "Augmented IDs",
        "Final gallery",
        "Nominal membership",
        "Operational reachability",
    ]
    for j, t in enumerate(h2):
        tbl2.rows[0].cells[j].text = t
    hv, iv = reach["hvidb"], reach["intact_cross_test"]
    rows = [
        [
            "HVIDB test",
            str(hv["nominal_biological_universe_ids"]),
            str(hv["sequence_available_retrievable_gallery"]),
            str(hv["augmentation_added_ids"]),
            str(hv["final_stated_evaluation_gallery"]),
            pct(hv["nominal_membership"]),
            pct(hv["operational_reachability"]),
        ],
        [
            "IntAct cross-test",
            str(iv["nominal_biological_universe_ids"]),
            str(iv["sequence_available_from_hvidb_nominal"]),
            str(iv["augmentation_added_non_nominal_true_partner_ids"]),
            str(iv["final_stated_evaluation_gallery"]),
            pct(iv["nominal_membership"]),
            pct(iv["operational_reachability"]),
        ],
    ]
    for vals in rows:
        r = tbl2.add_row()
        for j, v in enumerate(vals):
            r.cells[j].text = v
    for row in tbl2.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                set_line(p)
                set_run_tnr(p, "18")

    add(
        "Note. "
        + reach["interpretation"]
        + " IntAct augmentation adds 114 non-nominal true-partner viruses so ranking remains "
        "defined for every query; nominal membership (34.8%) reports coverage of the "
        "HVIDB-2104 nominal universe separately."
    )

    # S7 wording in supp if present
    for p in doc.paragraphs:
        if "purposive illustrative snapshot" in p.text:
            set_paragraph_text(
                p, p.text.replace("purposive illustrative snapshot", "illustrative audit")
            )
            set_line(p)
            set_run_tnr(p)


def main() -> None:
    # ensure analysis jsons
    import subprocess, sys

    subprocess.check_call(
        [sys.executable, str(ROOT / "scripts/compute_multilabel_any_known_positive.py")],
        cwd=str(ROOT),
    )
    subprocess.check_call(
        [sys.executable, str(ROOT / "scripts/compute_gallery_construction_reachability.py")],
        cwd=str(ROOT),
    )
    multi = json.loads(MULTI.read_text())
    reach = json.loads(REACH.read_text())

    main_doc = Document(str(MAIN))
    patch_main(main_doc, multi, reach)
    # dedupe reachability sentence in Table 2 note if doubled
    for p in main_doc.paragraphs:
        if p.text.startswith("Note. Ranking instantiations"):
            t = p.text
            phrase = (
                " Reachability was 100% by evaluation-gallery construction/filtering "
                "(operational, not natural database coverage; Supplementary Table S19)."
            )
            # older phrase
            old = (
                " Reachability was 100% by evaluation-gallery construction and should not be "
                "interpreted as empirical retrieval success"
            )
            if phrase in t and "should not be interpreted as empirical retrieval success" in t:
                # keep one clear sentence
                t2 = t
                # remove older long clause start if both
                pass
            if t.count("Supplementary Table S19") > 1:
                # leave
                pass
            break
    wc_with, wc_without = word_counts(main_doc)
    main_doc.save(str(MAIN))

    supp = Document(str(SUPP))
    append_supp(supp, multi, reach)
    supp.save(str(SUPP))

    update_cover(wc_with, wc_without)

    readme = ROOT / "README_转投说明.md"
    if readme.exists():
        text = readme.read_text()
        stamp = (
            f"**2026-08-27 更新（澄清/可复现）：** S18 any-known-positive 多标签敏感性；"
            f"S19 gallery construction / operational reachability；GAT 训练-only 澄清；"
            f"S7 illustrative audit。Cover letter 字数 **{wc_with:,} / {wc_without:,}**。\n\n"
        )
        if "澄清/可复现" not in text:
            lines = text.splitlines(keepends=True)
            insert_at = 1
            for i, line in enumerate(lines[:20]):
                if line.startswith("**2026"):
                    insert_at = i
                    break
            lines.insert(insert_at, stamp)
            readme.write_text("".join(lines))

    checklist = ROOT / "SUBMISSION_CHECKLIST.md"
    if checklist.exists():
        c = checklist.read_text().replace("Table S1–S17", "Table S1–S19")
        checklist.write_text(c)

    print({"wc_with": wc_with, "wc_without": wc_without})


if __name__ == "__main__":
    main()
