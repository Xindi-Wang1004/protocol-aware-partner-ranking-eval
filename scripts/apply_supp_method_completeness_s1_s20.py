#!/usr/bin/env python3
"""Supplementary method completeness edits (S1–S16 wording + S14 layout + S20)."""
from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

ROOT = Path(__file__).resolve().parents[1]
SUPP = ROOT / "02_Supplementary_Material.docx"
MAIN = ROOT / "01_Main_Manuscript.docx"
STRICT = ROOT / "analysis" / "intact_strict_nominal_gallery_sensitivity.json"

LINE, RULE = "480", "auto"


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = text


def set_cell(cell, text: str) -> None:
    if not cell.paragraphs:
        cell.text = text
        return
    set_paragraph_text(cell.paragraphs[0], text)
    for p in cell.paragraphs[1:]:
        p.text = ""


def ensure_pPr(p):
    pPr = p._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = etree.Element(qn("w:pPr"))
        p._element.insert(0, pPr)
    return pPr


def set_line(p):
    pPr = ensure_pPr(p)
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = etree.SubElement(pPr, qn("w:spacing"))
    sp.set(qn("w:line"), LINE)
    sp.set(qn("w:lineRule"), RULE)


def set_run_tnr(p, half="22"):
    for run in p.runs:
        rPr = run._element.find(qn("w:rPr"))
        if rPr is None:
            rPr = etree.SubElement(run._element, qn("w:rPr"))
            run._element.insert(0, rPr)
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = etree.SubElement(rPr, qn("w:rFonts"))
            rPr.insert(0, rFonts)
        for a in (qn("w:asciiTheme"), qn("w:hAnsiTheme"), qn("w:eastAsiaTheme"), qn("w:cstheme")):
            if a in rFonts.attrib:
                del rFonts.attrib[a]
        for a in (qn("w:ascii"), qn("w:hAnsi"), qn("w:cs"), qn("w:eastAsia")):
            rFonts.set(a, "Times New Roman")
        for tag in ("w:sz", "w:szCs"):
            el = rPr.find(qn(tag))
            if el is None:
                el = etree.SubElement(rPr, qn(tag))
            el.set(qn("w:val"), half)


def insert_after(paragraph, text: str):
    new_p = paragraph._element.makeelement(qn("w:p"), {})
    paragraph._element.addnext(new_p)
    new_para = type(paragraph)(new_p, paragraph._parent)
    new_para.add_run(text)
    set_line(new_para)
    set_run_tnr(new_para)
    return new_para


def insert_before(paragraph, text: str):
    new_p = paragraph._element.makeelement(qn("w:p"), {})
    paragraph._element.addprevious(new_p)
    new_para = type(paragraph)(new_p, paragraph._parent)
    new_para.add_run(text)
    set_line(new_para)
    set_run_tnr(new_para)
    return new_para


def pct(x: float, d: int = 1) -> str:
    return f"{100 * x:.{d}f}%"


def main() -> None:
    strict = json.loads(STRICT.read_text())
    s = Document(str(SUPP))

    # --- 1. S13 Oracle@10 → Oracle@20 Hit@10 ---
    for table in s.tables:
        hdr = table.rows[0].cells[-1].text.strip()
        if hdr == "Oracle@10" or hdr.replace("\n", " ") == "Oracle@10":
            set_cell(table.rows[0].cells[-1], "Oracle@20 Hit@10")
            print("fixed S13 Oracle column")

    # --- 2. Move S14 table (Role/Source...) next to S14 caption/note ---
    s14_note = None
    for p in s.paragraphs:
        if p.text.startswith("Note.") and "Pair counts include constructed negatives" in p.text:
            s14_note = p
            break
    role_table = None
    for table in s.tables:
        if table.rows[0].cells[0].text.strip() == "Role" and "Source" in table.rows[0].cells[1].text:
            role_table = table
            break
    if s14_note is not None and role_table is not None:
        tbl_el = role_table._tbl
        parent = tbl_el.getparent()
        parent.remove(tbl_el)
        s14_note._element.addnext(tbl_el)
        print("moved S14 table after its Note")
    else:
        print("WARN: S14 move skipped", s14_note is not None, role_table is not None)

    # Reload paragraph indices after move? paragraph list is fine; table moved in XML.
    # Re-open to refresh python-docx caches for later appends
    s.save(str(SUPP))
    s = Document(str(SUPP))

    # --- 3. S6 MRR miss handling ---
    for p in s.paragraphs:
        if p.text.startswith("Note.") and "Mean reciprocal rank (MRR)" in p.text:
            if "retrieval miss is assigned reciprocal rank 0" not in p.text:
                set_paragraph_text(
                    p,
                    p.text.rstrip(".")
                    + ". For shortlist-based fusion, a retrieval miss is assigned reciprocal rank 0 "
                    "in the full-query MRR calculation.",
                )
                print("S6 MRR miss note")
            break

    # --- 4. S16 no-hit handling ---
    for p in s.paragraphs:
        if p.text.startswith("Note.") and "MMseqs2 easy-search" in p.text:
            if "queries without reported hits" not in p.text:
                set_paragraph_text(
                    p,
                    p.text.rstrip(".")
                    + " MMseqs2 was used only to generate a score-based ordering; queries without "
                    "reported hits were assigned score 0 for all unreported gallery members "
                    "(bottom-tied under bitscore ranking), with deterministic tie handling by "
                    "stable gallery index order.",
                )
                print("S16 no-hit note")
            break

    # --- 5. Soften S4 wording ---
    for p in s.paragraphs:
        if "Pair-subset corresponds to the smallest nested sets in this controlled curve" in p.text:
            set_paragraph_text(
                p,
                p.text.replace(
                    "Pair-subset corresponds to the smallest nested sets in this controlled curve.",
                    "The label-informed pair-subset protocol behaves analogously to the small-gallery "
                    "regime in this controlled simulation, although its per-query galleries are not "
                    "themselves a single globally nested sequence.",
                ),
            )
            print("S4 wording")
            break

    # --- 6. Scope and protocol exceptions (after Protocol notes heading) ---
    for i, p in enumerate(s.paragraphs):
        if p.text.strip() == "Protocol notes":
            # Insert before the next heading if not already present
            already = any(
                "Scope and protocol exceptions" in q.text for q in s.paragraphs[i : i + 8]
            )
            if not already:
                scope_title = insert_after(p, "Scope and protocol exceptions")
                scope_body = (
                    "Unless noted, fixed-gallery diagnostics use a dataset-specific stated evaluation "
                    "gallery that is fixed before scoring for every query in that evaluation file. "
                    "“Fixed” therefore means a shared, pre-specified candidate set within an evaluation, "
                    "not that the same viral ID list is reused unchanged across datasets. The HVIDB-2104 "
                    "nominal ID universe is shared as the nominal biological reference, but the "
                    "retrievable / stated galleries may differ by sequence availability and evaluation "
                    "context. For the IntAct cross-test, 114 true-partner viruses outside HVIDB-2104 "
                    "were added so that ranking remains defined for every labelled query (operational "
                    "reachability 100%); natural coverage of the HVIDB-2104 nominal universe is reported "
                    "separately as nominal membership (34.8%). A strict HVIDB-2104 sequence-available "
                    "gallery without that augmentation is reported as a sensitivity analysis "
                    "(Supplementary Table S20)."
                )
                insert_after(scope_title, scope_body)
                print("added Scope and protocol exceptions")
            break

    # --- 7. Metric definitions: three reachability notions ---
    for p in s.paragraphs:
        if p.text.startswith("Retrievable-partner reachability:"):
            set_paragraph_text(
                p,
                "Strict sequence reachability: fraction of queries whose true viral partner has an "
                "available sequence and is present in the sequence-available retrievable gallery "
                "derived from the HVIDB-2104 nominal universe (no cross-dataset augmentation).",
            )
            print("rewrote reachability bullet 1")
        elif p.text.startswith("Nominal-gallery membership:"):
            set_paragraph_text(
                p,
                "Nominal membership: fraction of queries whose labelled true partner belongs to the "
                "HVIDB-2104 nominal biological ID set (identifier membership, independent of whether "
                "a sequence is available for scoring).",
            )
        elif p.text.startswith("Bootstrap 95% CI:"):
            # insert operational reachability before bootstrap if missing
            if not any(
                q.text.startswith("Operational reachability:") for q in s.paragraphs
            ):
                insert_before(
                    p,
                    "Operational reachability: fraction of queries whose true partner is eligible in "
                    "the stated evaluation gallery actually used for ranking (after any declared "
                    "augmentation/filtering). This is the Reachability factor in the main-text "
                    "identity Hit@K = Reachability × Recall@T|reachable × oracle@T Hit@K.",
                )
                print("added operational reachability bullet")
            break

    # --- 8–9. S2 Sigmoid/logit + LoRA ---
    for p in s.paragraphs:
        if p.text.startswith("• Stage-1 LoRA:"):
            set_paragraph_text(
                p,
                "• Stage-1 LoRA: LoRA adapters (96 Linear layers; rank=8, α_LoRA=16, dropout=0) were "
                "retained in the architecture for upstream code compatibility but were never trained. "
                "At the reported Stage-1 checkpoint they remain at random initialisation and contribute "
                "no learned adaptation, so Stage-1 is effectively frozen-ESM3 retrieval.",
            )
            print("S2 LoRA")
        elif p.text.startswith("• Stage-2 classifier:"):
            set_paragraph_text(
                p,
                "• Stage-2 classifier: frozen ESM3 + MLP 3072→256→128→1 (ReLU, dropout 0.1); a final "
                "Sigmoid is applied for BCE training on balanced positive/negative pairs (~1:1; random + "
                "mixed hard negatives), whereas ranking and score fusion use the pre-Sigmoid logit; "
                "AdamW lr=2×10⁻⁵; batch size=1; up to 10 epochs; model selection by best HVIDB "
                "validation F1 (checkpoint esm3_frozen).",
            )
            print("S2 Sigmoid/logit")

    # --- 10. S1 negatives not in gallery ---
    for p in s.paragraphs:
        if p.text.startswith("Note.") and "The IntAct cross-test used in the main text" in p.text:
            if "defined over viral IDs rather than the constructed pair file" not in p.text:
                set_paragraph_text(
                    p,
                    p.text.rstrip(".")
                    + " Constructed negatives are not added to the fixed evaluation gallery because "
                    "the fixed-gallery protocol is defined over viral IDs rather than the constructed "
                    "pair file.",
                )
                print("S1 negatives note")
            break

    # --- 11. S7 coding note ---
    for p in s.paragraphs:
        if p.text.startswith("Note.") and "Illustrative audit of n=12" in p.text:
            if "without independent dual review" not in p.text:
                set_paragraph_text(
                    p,
                    p.text.rstrip(".")
                    + " Source articles were accessed from publisher/PDF pages during manuscript "
                    "preparation (2025–2026). Coding was performed by one author using the checklist "
                    "above and was not independently dual-reviewed.",
                )
                print("S7 coding note")
            break

    # --- 12. S5 35-fold wording ---
    for p in s.paragraphs:
        if p.text.startswith("Note.") and "~35-fold HVIDB shift" in p.text:
            t = p.text.replace(
                "The ~35-fold HVIDB shift (91.3%→2.6%) is the control cited in the main text; it decomposes approximately into candidate-set composition (91.3%/7.3%≈12.5 using direct full-gallery classifier ranking as the no-gate fixed-gallery comparator) and Stage-1 gating (7.3%/2.6%≈2.8).",
                "The HVIDB contrast is a 35.1-fold reduction in Hit@10, from 91.3% to 2.6%, as cited in the main text; it decomposes approximately into candidate-set composition (91.3%/7.3%≈12.5 using direct full-gallery classifier ranking as the no-gate fixed-gallery comparator) and Stage-1 gating (7.3%/2.6%≈2.8). These factors are descriptive diagnostics rather than independent causal estimates.",
            )
            set_paragraph_text(p, t)
            print("S5 35.1-fold wording")
            break

    # --- 13. S20 strict IntAct gallery ---
    if not any(
        p.text.startswith("Supplementary Table S20.") for p in s.paragraphs
    ):
        aug = strict["augmented_operational"]
        st = strict["strict_nominal_sequence_gallery"]
        # append at end
        cap = s.add_paragraph(
            "Supplementary Table S20. IntAct cross-test: augmented operational gallery versus "
            "strict HVIDB-2104 sequence-available gallery."
        )
        set_line(cap)
        set_run_tnr(cap)
        tbl = s.add_table(rows=1, cols=6)
        headers = [
            "Gallery protocol",
            "Reachability",
            "Recall@20",
            "Recall@100",
            "Fusion Hit@10",
            "Oracle@20 Hit@10 (n)",
        ]
        for j, h in enumerate(headers):
            tbl.rows[0].cells[j].text = h
        rows = [
            [
                "Augmented (main; +114 non-nominal)",
                pct(aug["reachability"]),
                pct(aug["recall@20"]),
                pct(aug["recall@100"]),
                pct(aug["fusion_hit@10"]),
                "— (see Table 2)",
            ],
            [
                "Strict HVIDB-2104 sequence-available (n=1,115)",
                pct(st["reachability"]),
                pct(st["recall@20_all_queries_unreachable_as_miss"]),
                pct(st["recall@100_all_queries_unreachable_as_miss"]),
                pct(st["fusion_hit@10_all_queries"]),
                f"{pct(st['oracle@20_fusion_hit@10'])} (n={st['oracle@20_n_reachable']})",
            ],
        ]
        for row in rows:
            r = tbl.add_row()
            for j, v in enumerate(row):
                r.cells[j].text = v
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    set_line(p)
                    set_run_tnr(p, "18")
        note = s.add_paragraph(
            "Note. Strict gallery = the 1,115 sequence-available HVIDB-2104 viruses with no IntAct "
            f"non-nominal augmentation (114 IDs excluded). Natural/strict reachability = "
            f"{pct(st['reachability'])} (n={strict['n_reachable']}/{strict['n_queries']}); "
            "unreachable queries contribute 0 to end-to-end Hit@K and are treated as retrieval misses "
            "in the all-query Recall@K columns above. Reachable-only rates under the strict definition: "
            f"Recall@20 {pct(st['recall@20_reachable_only'])}, Recall@100 {pct(st['recall@100_reachable_only'])}, "
            f"fusion Hit@10 {pct(st['fusion_hit@10_reachable_only'])}. Per-query Stage-1/Stage-2 indicators "
            "for reachable queries are taken from the deposited augmented-gallery evaluation; removing "
            "the 114 non-nominal decoys can only improve or leave unchanged a reachable query’s rank, so "
            "reachable-subset rates are conservative lower bounds under the strict gallery. Identity "
            "residual for Reachability × Recall@20|reachable × oracle@20 Hit@10 versus end-to-end "
            f"fusion Hit@10 is {st['identity_abs_error_vs_reach_x_R20_x_oracle']:.2e}."
        )
        set_line(note)
        set_run_tnr(note)
        print("added S20")

    s.save(str(SUPP))

    # --- Main: operational Reachability in formula prose (minimal) ---
    m = Document(str(MAIN))
    for p in m.paragraphs:
        if p.text.startswith(
            "Under the constrained retrieve-then-rerank protocol studied here, headline end-to-end Hit@K"
        ):
            if "operational reachability" not in p.text:
                set_paragraph_text(
                    p,
                    p.text.replace(
                        "Let Reachability be the fraction of labelled evaluation queries whose true partner is eligible in the stated evaluation gallery.",
                        "Let Reachability denote operational reachability—the fraction of labelled evaluation queries whose true partner is eligible in the stated evaluation gallery actually used for ranking (Supplementary Protocol notes; Supplementary Tables S19–S20).",
                    ),
                )
                print("main operational Reachability")
            break
    m.save(str(MAIN))
    print("done")


if __name__ == "__main__":
    main()
